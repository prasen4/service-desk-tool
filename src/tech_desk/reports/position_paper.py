"""Per-vendor Position Paper generation.

Two-stage LLM pipeline modeled on the Cotiviti NotebookLM -> ChatGPT workflow:

1. Research brief — gathers internal context (CRM notes, curated updates,
   attachment excerpts) plus a fresh web search, then asks the LLM to distill
   it into a structured research brief (JSON).
2. Position paper — takes the research brief (+ optional analyst custom
   prompt) and asks the LLM to draft the full position paper sections, which
   are then rendered into a .docx matching the standard template structure.
"""

from __future__ import annotations

import json
import logging
import re
from pathlib import Path

from tech_desk.config import get_settings
from tech_desk.database import PositionPaperORM, get_session_factory, init_db, position_paper_from_orm
from tech_desk.llm import LLMClient
from tech_desk.models import ComparisonRow, CompetitorProfile, PositionPaperResult
from tech_desk.research.search import WebSearcher
from tech_desk.timeutils import now_utc

logger = logging.getLogger(__name__)

RESEARCH_BRIEF_SYSTEM_PROMPT = """You are a research analyst preparing a briefing document about a
technology vendor, for Cotiviti (a healthcare payer analytics and payment integrity company). Synthesize
the provided internal notes, prior curated updates, and fresh web search results into a research brief,
but do not limit yourself to only what's in that material -- you also have broad general knowledge about
this vendor, its product category, and its market. Use that general knowledge freely to fill gaps (e.g.
naming real, well-known competitors and describing typical pricing/positioning for this category),
especially when the supplied evidence is thin. Clearly distinguish sourced claims (cite the source URL)
from general-knowledge claims (mark them as such) -- but never leave a section empty just because web
search evidence was sparse. It is far more useful to provide an informed, clearly-labeled general-knowledge
answer than an empty one."""

POSITION_PAPER_SYSTEM_PROMPT = """You are drafting a Vendor Position Paper for Cotiviti technology
leadership, in the style of a professional vendor/market evaluation (comparable to how an industry
analyst profiles a SaaS or AI vendor for a prospective enterprise buyer). Write a genuinely general-purpose
position paper about what this vendor's product/service does and how it's positioned in its market -- then
connect that to Cotiviti's healthcare payer analytics / payment integrity domain specifically in the
"Cotiviti Relevance" section.

Ground claims in the supplied research brief where possible, but you are expected and encouraged to also
draw on your own general knowledge of this vendor, its product category, and its known competitors -- the
brief is a starting point, not a hard ceiling. Always name specific, real competitor products (never leave
the competitors list or comparison table empty) -- if a claim comes from general knowledge rather than the
brief's sourced evidence, set "source_verified": false on that item, but still include it."""


def _slugify(name: str) -> str:
    slug = re.sub(r"[^a-zA-Z0-9]+", "-", name.strip().lower()).strip("-")
    return slug or "vendor"


def _gather_internal_context(session, vendor_name: str) -> str:
    """Pull CRM notes, status, attachment excerpts, and recent curated updates
    for this vendor into a single text block for the LLM prompt."""
    from tech_desk import vendor_profiles
    from tech_desk.vendors import get_vendor_updates

    parts: list[str] = []

    profile = vendor_profiles.get_vendor_profile(session, vendor_name)
    if profile:
        parts.append(f"CRM status: {profile.get('status_label') or 'Not tracked in CRM'}")
        notes = profile.get("notes") or []
        if notes:
            parts.append("Analyst notes:")
            for n in notes[:8]:
                author = f" ({n['author']})" if n.get("author") else ""
                parts.append(f"- [{n['created_at'][:10]}]{author} {n['body']}")
        attachments = profile.get("attachments") or []
        text_attachments = [a for a in attachments if a["filename"].lower().endswith(".txt")]
        if text_attachments:
            parts.append("Attachment excerpts:")
            for a in text_attachments[:5]:
                attachment_orm = vendor_profiles.get_attachment(session, vendor_name, a["id"])
                if attachment_orm is None:
                    continue
                path = vendor_profiles.attachment_path(attachment_orm)
                if path.exists():
                    try:
                        excerpt = path.read_text(encoding="utf-8", errors="ignore")[:1500]
                        parts.append(f"- {a['filename']}: {excerpt}")
                    except OSError:
                        pass

    news = get_vendor_updates(session, vendor_name, limit=15)
    if news and news.get("updates"):
        parts.append("Recent curated updates (from prior research runs):")
        for u in news["updates"][:15]:
            parts.append(f"- [{u.get('published_date') or u.get('discovered_at', '')[:10]}] {u['title']} — {u['summary']} ({u['source_url']})")

    return "\n".join(parts) if parts else "No internal CRM notes or prior research updates on file for this vendor."


def _run_web_search(searcher: WebSearcher, vendor_name: str) -> str:
    """Fresh, targeted web search for this vendor — supplements (does not
    replace) whatever internal context already exists."""
    queries = [
        f"{vendor_name} healthcare technology overview",
        f"{vendor_name} pricing licensing model",
        f"{vendor_name} competitors alternatives comparison",
    ]
    seen_urls: set[str] = set()
    lines: list[str] = []
    for query in queries:
        try:
            results = searcher.search(query, timelimit="y")
        except Exception as exc:
            logger.warning("Web search failed for '%s': %s", query, exc)
            continue
        for r in results[:5]:
            if r.url in seen_urls:
                continue
            seen_urls.add(r.url)
            lines.append(f"- {r.title} — {r.snippet} ({r.url})")
    return "\n".join(lines) if lines else "No additional web search results found."


class PositionPaperGenerator:
    def __init__(self, llm: LLMClient | None = None):
        self.llm = llm or LLMClient()
        self.searcher = WebSearcher()

    def generate(
        self,
        vendor_name: str,
        *,
        custom_prompt: str = "",
        progress=None,
    ) -> PositionPaperResult:
        def _p(msg: str, pct: int) -> None:
            if progress:
                progress(msg, pct)

        init_db()
        session = get_session_factory()()
        orm = PositionPaperORM(vendor=vendor_name, status="running", custom_prompt=(custom_prompt or "").strip())
        session.add(orm)
        session.commit()
        session.refresh(orm)

        try:
            _p("Gathering internal CRM context...", 10)
            internal_context = _gather_internal_context(session, vendor_name)

            _p("Running fresh web search...", 25)
            web_findings = _run_web_search(self.searcher, vendor_name)

            _p("Building research brief...", 40)
            brief = self._build_research_brief(vendor_name, internal_context, web_findings)
            orm.research_brief_json = json.dumps(brief)
            session.commit()

            _p("Drafting position paper...", 60)
            paper = self._build_position_paper(vendor_name, brief, custom_prompt)

            _p("Rendering document...", 85)
            docx_path = self._render_docx(vendor_name, paper)

            orm.status = "completed"
            orm.docx_path = str(docx_path)
            orm.generated_at = now_utc()
            session.commit()
            session.refresh(orm)

            _p("Position paper ready", 100)
            return position_paper_from_orm(orm)

        except Exception as exc:
            orm.status = "failed"
            orm.error_message = str(exc)
            session.commit()
            logger.exception("Position paper generation failed for %s", vendor_name)
            raise
        finally:
            session.close()

    def _build_research_brief(self, vendor_name: str, internal_context: str, web_findings: str) -> dict:
        prompt = f"""Vendor: {vendor_name}

Internal context (Cotiviti CRM notes + prior curated research):
{internal_context}

Fresh web search findings:
{web_findings}

Respond in JSON:
{{
  "snapshot": "2-3 sentences: what the vendor does and its market position",
  "capabilities": ["capability 1", "capability 2"],
  "differentiators": ["what sets them apart"],
  "limitations_risks": ["known limitations, risks, or open questions"],
  "integration_implementation": "notes on integration/implementation approach (use general knowledge of this product category if the evidence is thin)",
  "security_compliance": "notes on security/compliance posture (HIPAA, SOC2, etc.) — use general knowledge of typical posture for this vendor/category if not explicitly found",
  "pricing_licensing": "notes on pricing/licensing model — use general knowledge of typical pricing for this category if not explicitly found",
  "cotiviti_relevance": "why this matters specifically for Cotiviti's payer analytics / payment integrity mission",
  "competitors": [{{"name": "", "why_relevant": "", "key_differences": "", "evidence": "", "source_verified": true}}],
  "references": [{{"label": "", "type": "Vendor Documentation|Review Article|News Article|Analyst Report|Other", "url": ""}}]
}}

IMPORTANT: "competitors" must always include at least 2-3 real, specifically-named products that compete with
{vendor_name} in its category — draw on general market knowledge for this if the web findings above didn't
surface any by name. Never return an empty competitors list."""
        try:
            # Reasoning models (e.g. gpt-5.4) spend part of this budget on
            # hidden reasoning tokens before emitting visible output, so this
            # needs real headroom above what the JSON schema itself requires —
            # too low a limit here previously produced truncated/empty JSON.
            return self.llm.chat_json(RESEARCH_BRIEF_SYSTEM_PROMPT, prompt, temperature=0.3, max_tokens=6000)
        except Exception as exc:
            logger.warning("Research brief generation failed for %s: %s", vendor_name, exc)
            return {"snapshot": f"Research brief unavailable: {exc}"}

    def _build_position_paper(self, vendor_name: str, brief: dict, custom_prompt: str) -> dict:
        instructions_block = f"\nAdditional analyst guidance for this paper: {custom_prompt.strip()}\n" if custom_prompt.strip() else ""
        prompt = f"""Vendor: {vendor_name}

Research brief (starting point — you may also draw on your own general knowledge, per your instructions):
{json.dumps(brief, indent=2)}
{instructions_block}
Draft the full Position Paper. Respond in JSON:
{{
  "executive_summary": "1-2 paragraphs: what {vendor_name} does, its market position, and why it's relevant to organizations operationalizing this category of technology",
  "capabilities": ["bullet 1", "bullet 2"],
  "differentiators": ["bullet 1", "bullet 2"],
  "limitations_risks": ["bullet 1", "bullet 2"],
  "integration_implementation": "1-2 paragraphs",
  "security_compliance": "1 paragraph",
  "pricing_licensing": "1 paragraph",
  "cotiviti_relevance": "1-2 paragraphs on implications specifically for Cotiviti's payer analytics / payment integrity mission",
  "competitors": [{{"name": "", "why_relevant": "", "key_differences": "", "evidence": "", "source_verified": true}}],
  "comparison_table": [
    {{"criterion": "Integration with current systems", "values": {{"{vendor_name}": "...", "<Competitor A name>": "...", "<Competitor B name>": "..."}}, "notes": ""}},
    {{"criterion": "Cost / Licensing Model", "values": {{"{vendor_name}": "...", "<Competitor A name>": "...", "<Competitor B name>": "..."}}, "notes": ""}},
    {{"criterion": "Security / Compliance", "values": {{"{vendor_name}": "...", "<Competitor A name>": "...", "<Competitor B name>": "..."}}, "notes": ""}},
    {{"criterion": "Scalability / Performance", "values": {{"{vendor_name}": "...", "<Competitor A name>": "...", "<Competitor B name>": "..."}}, "notes": ""}},
    {{"criterion": "Overall Fit", "values": {{"{vendor_name}": "...", "<Competitor A name>": "...", "<Competitor B name>": "..."}}, "notes": ""}}
  ],
  "architecture_overview": "1-2 sentences describing how this product typically fits into a customer's technology stack",
  "key_insights": ["3-5 short bullets: key insights & implications for a healthcare payer analytics buyer"],
  "recommendation": "1 short paragraph: analyst take / recommendation on how/when to consider this vendor",
  "references": [{{"label": "", "type": "Vendor Documentation|Review Article|News Article|Analyst Report|Other", "url": ""}}]
}}

IMPORTANT: use the SAME real competitor product names (as keys in "values") consistently across every row of
"comparison_table", and use those same names in the "competitors" list. Never return an empty competitors list
or an empty comparison_table — use general knowledge of this vendor's market to name 2-3 real competitors if the
research brief didn't supply any."""
        try:
            # Same reasoning-token headroom concern as the research brief —
            # this prompt asks for a longer, multi-section document (exec
            # summary, comparison table, insights, etc.), so it needs even more.
            return self.llm.chat_json(POSITION_PAPER_SYSTEM_PROMPT, prompt, temperature=0.4, max_tokens=8000)
        except Exception as exc:
            logger.warning("Position paper draft failed for %s: %s", vendor_name, exc)
            return {"executive_summary": f"Position paper draft unavailable: {exc}"}

    def _render_docx(self, vendor_name: str, paper: dict) -> Path:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Inches, Pt, RGBColor

        settings = get_settings()
        out_dir = settings.tech_desk_data_dir / "exports" / "position_papers"
        out_dir.mkdir(parents=True, exist_ok=True)
        stamp = now_utc().strftime("%Y%m%d_%H%M%S")
        out_path = out_dir / f"{_slugify(vendor_name)}_{stamp}.docx"

        # Brand palette + typography lifted directly from Cotiviti's official
        # "Position Paper Template" (.odt) — extracted field-by-field so the
        # generated .docx matches its cover layout, colors, fonts, header/
        # footer, and table styling.
        PURPLE = RGBColor(0x30, 0x00, 0x6F)
        PINK = RGBColor(0xEC, 0x00, 0x8C)
        GRAY = RGBColor(0x5E, 0x61, 0x79)
        BORDER_GRAY = "A6A6A6"

        assets_dir = Path(__file__).resolve().parent / "assets"
        logo_path = assets_dir / "cotiviti_logo.png"
        banner_path = assets_dir / "cover_banner.png"

        def add_floating_picture(
            paragraph, image_path: Path, width_in: float, height_in: float | None, x_in: float, y_in: float, behind: bool = True
        ) -> None:
            """Insert a page-anchored (floating) picture — used for the cover
            page's full-height banner graphic and logo, matching the template's
            header artwork rather than an inline image that would push text.

            height_in may be None, in which case python-docx computes it from
            the image's own native aspect ratio (avoids stretching/squashing
            logos that don't exactly match a hardcoded width:height ratio).
            """
            if not image_path.exists():
                return
            run = paragraph.add_run()
            if height_in is None:
                run.add_picture(str(image_path), width=Inches(width_in))
            else:
                run.add_picture(str(image_path), width=Inches(width_in), height=Inches(height_in))
            drawing = run._r.find(qn("w:drawing"))
            inline = drawing.find(qn("wp:inline"))
            extent = inline.find(qn("wp:extent"))
            doc_pr = inline.find(qn("wp:docPr"))
            frame_locks = inline.find(qn("wp:cNvGraphicFramePr"))
            graphic = inline.find(qn("a:graphic"))

            anchor = OxmlElement("wp:anchor")
            for attr, val in {
                "distT": "0", "distB": "0", "distL": "0", "distR": "0",
                "simplePos": "0", "relativeHeight": "251659264",
                "behindDoc": "1" if behind else "0", "locked": "0",
                "layoutInCell": "1", "allowOverlap": "1",
            }.items():
                anchor.set(attr, val)

            simple_pos = OxmlElement("wp:simplePos")
            simple_pos.set("x", "0")
            simple_pos.set("y", "0")
            anchor.append(simple_pos)

            pos_h = OxmlElement("wp:positionH")
            pos_h.set("relativeFrom", "page")
            pos_h_off = OxmlElement("wp:posOffset")
            pos_h_off.text = str(int(Inches(x_in)))
            pos_h.append(pos_h_off)
            anchor.append(pos_h)

            pos_v = OxmlElement("wp:positionV")
            pos_v.set("relativeFrom", "page")
            pos_v_off = OxmlElement("wp:posOffset")
            pos_v_off.text = str(int(Inches(y_in)))
            pos_v.append(pos_v_off)
            anchor.append(pos_v)

            anchor.append(extent)
            effect_extent = OxmlElement("wp:effectExtent")
            for attr in ("l", "t", "r", "b"):
                effect_extent.set(attr, "0")
            anchor.append(effect_extent)
            anchor.append(OxmlElement("wp:wrapNone"))
            anchor.append(doc_pr)
            anchor.append(frame_locks)
            anchor.append(graphic)

            drawing.remove(inline)
            drawing.append(anchor)

        def add_page_number_field(paragraph) -> None:
            run = paragraph.add_run()
            fld_begin = OxmlElement("w:fldChar")
            fld_begin.set(qn("w:fldCharType"), "begin")
            instr = OxmlElement("w:instrText")
            instr.set(qn("xml:space"), "preserve")
            instr.text = "PAGE"
            fld_sep = OxmlElement("w:fldChar")
            fld_sep.set(qn("w:fldCharType"), "separate")
            fld_end = OxmlElement("w:fldChar")
            fld_end.set(qn("w:fldCharType"), "end")
            for el in (fld_begin, instr, fld_sep, fld_end):
                run._r.append(el)
            run.font.size = Pt(8)
            run.font.color.rgb = GRAY

        def shade_cell(cell, hex_color: str) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), hex_color)
            tc_pr.append(shd)

        def border_cell(cell, hex_color: str = BORDER_GRAY) -> None:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                el = OxmlElement(f"w:{edge}")
                el.set(qn("w:val"), "single")
                el.set(qn("w:sz"), "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), hex_color)
                borders.append(el)
            tc_pr.append(borders)

        def set_cell(cell, text: str, *, header: bool = False, center: bool = False) -> None:
            cell.text = ""
            p = cell.paragraphs[0]
            if center or header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.font.size = Pt(10)
            if header:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                shade_cell(cell, PURPLE.__str__())
            border_cell(cell)

        doc = Document()

        # --- Page geometry (US Letter, matches the template's margins) ---
        section = doc.sections[0]
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        base_font = doc.styles["Normal"].font
        base_font.name = "Calibri"
        base_font.size = Pt(11)

        # --- Cover page has its own header (logo + banner graphic, no
        # footer); body pages get a small header logo + copyright footer —
        # mirrors the template's two master pages (MP0 cover / MP1 body). ---
        section.different_first_page_header_footer = True

        fp_header = section.first_page_header
        fp_header.is_linked_to_previous = False
        fp_header_p = fp_header.paragraphs[0]
        fp_header_p.text = ""
        add_floating_picture(fp_header_p, logo_path, 2.36, None, 1.0, 0.75, behind=False)
        add_floating_picture(fp_header_p, banner_path, 2.84, 8.63, 5.66, 0.0, behind=True)

        fp_footer = section.first_page_footer
        fp_footer.is_linked_to_previous = False
        fp_footer.paragraphs[0].text = ""

        header = section.header
        header.is_linked_to_previous = False
        header_p = header.paragraphs[0]
        header_p.text = ""
        add_floating_picture(header_p, logo_path, 1.53, None, 5.97, 0.3, behind=False)

        footer = section.footer
        footer.is_linked_to_previous = False
        footer_p = footer.paragraphs[0]
        footer_p.text = ""
        footer_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5))
        copyright_run = footer_p.add_run(
            f"\u00a9 {now_utc().year} Cotiviti, Inc. All rights reserved. All proprietary information "
            "shall remain the sole and exclusive property of Cotiviti, Inc."
        )
        copyright_run.font.size = Pt(8)
        copyright_run.font.color.rgb = GRAY
        footer_p.add_run("\t")
        add_page_number_field(footer_p)

        # --- Cover page content (mirrors the template's title block, "For
        # the Delivery of" / contact panel, and classification/date lines) ---
        for _ in range(4):
            doc.add_paragraph()

        title_p = doc.add_paragraph()
        title_run = title_p.add_run(vendor_name)
        title_run.font.name = "Arial"
        title_run.font.size = Pt(36)
        title_run.font.bold = True
        title_run.font.color.rgb = PURPLE

        subtitle_p = doc.add_paragraph()
        subtitle_run = subtitle_p.add_run("Position Paper")
        subtitle_run.font.name = "Arial"
        subtitle_run.font.size = Pt(36)
        subtitle_run.font.color.rgb = PURPLE

        prepared_for_p = doc.add_paragraph()
        prepared_for_run = prepared_for_p.add_run("Prepared for: Cotiviti Technology Leadership")
        prepared_for_run.font.size = Pt(14)
        prepared_for_run.font.color.rgb = PINK

        for _ in range(3):
            doc.add_paragraph()

        delivery_p = doc.add_paragraph()
        delivery_run = delivery_p.add_run("For the Delivery of")
        delivery_run.font.size = Pt(12)
        delivery_run.font.color.rgb = PURPLE
        doc.add_paragraph("Cotiviti Enterprise AI R&D")

        for _ in range(2):
            doc.add_paragraph()

        contact_p = doc.add_paragraph()
        contact_run = contact_p.add_run("Contact")
        contact_run.font.size = Pt(12)
        contact_run.font.color.rgb = PURPLE

        for label in ("Name", "Title", "Cotiviti, Inc.", "Phone"):
            p = doc.add_paragraph()
            p.add_run(label).font.color.rgb = GRAY
        doc.add_paragraph("Email")

        for _ in range(2):
            doc.add_paragraph()
        doc.add_paragraph("Data Sensitivity Classification: Select from drop-down menu")
        doc.add_paragraph(now_utc().strftime("Prepared %B %d, %Y"))

        doc.add_page_break()

        def add_heading(text: str, level: int = 1) -> None:
            h = doc.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(text)
            if level == 1:
                run.font.size = Pt(20)
                run.font.bold = True
                run.font.color.rgb = PURPLE
            elif level == 2:
                run.font.size = Pt(14)
                run.font.color.rgb = PINK
            else:
                run.font.size = Pt(12)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)

        def add_bullets(items) -> None:
            if not items:
                doc.add_paragraph("No evidence available.")
                return
            for item in items:
                doc.add_paragraph(str(item), style="List Bullet")

        def add_body(text: str) -> None:
            doc.add_paragraph(text or "No evidence available.")

        def coerce_competitor(c) -> CompetitorProfile | None:
            """Best-effort coercion so a slightly-off-schema LLM response
            (e.g. a bare string, or a dict using a different name key)
            doesn't get silently dropped from the document."""
            if isinstance(c, str):
                return CompetitorProfile(name=c, source_verified=False)
            if isinstance(c, dict):
                name = c.get("name") or c.get("competitor") or c.get("vendor") or c.get("product")
                if not name:
                    return None
                try:
                    return CompetitorProfile(
                        name=str(name),
                        why_relevant=str(c.get("why_relevant", "")),
                        key_differences=str(c.get("key_differences", "")),
                        evidence=str(c.get("evidence", "")),
                        source_verified=bool(c.get("source_verified", False)),
                    )
                except Exception:
                    return None
            return None

        def coerce_comparison_row(r) -> ComparisonRow | None:
            if not isinstance(r, dict):
                return None
            criterion = r.get("criterion")
            if not criterion:
                return None
            values = r.get("values")
            if not isinstance(values, dict):
                # Tolerate the older fixed-column shape (product/competitor_1/competitor_2)
                values = {
                    k: str(r[k])
                    for k in ("product", "competitor_1", "competitor_2")
                    if r.get(k)
                }
            try:
                return ComparisonRow(
                    criterion=str(criterion),
                    values={str(k): str(v) for k, v in values.items()},
                    notes=str(r.get("notes", "")),
                )
            except Exception:
                return None

        add_heading("Executive Summary")
        add_body(paper.get("executive_summary") or paper.get("vendor_snapshot", ""))

        add_heading("Comparative Context")
        add_heading("Capabilities", level=2)
        add_bullets(paper.get("capabilities"))
        add_heading("Differentiators", level=2)
        add_bullets(paper.get("differentiators"))
        add_heading("Limitations & Risks", level=2)
        add_bullets(paper.get("limitations_risks"))

        add_heading("Integration & Implementation")
        add_body(paper.get("integration_implementation", ""))

        add_heading("Security & Compliance")
        add_body(paper.get("security_compliance", ""))

        add_heading("Pricing & Licensing")
        add_body(paper.get("pricing_licensing", ""))

        add_heading("Cotiviti Relevance")
        add_body(paper.get("cotiviti_relevance", ""))

        competitors = [c for c in (coerce_competitor(c) for c in (paper.get("competitors") or [])) if c]
        add_heading("Competitors")
        if competitors:
            for profile in competitors:
                p = doc.add_paragraph()
                p.add_run(profile.name + " ").bold = True
                if not profile.source_verified:
                    p.add_run("(general market knowledge)").italic = True
                if profile.why_relevant:
                    doc.add_paragraph(f"Why relevant: {profile.why_relevant}")
                if profile.key_differences:
                    doc.add_paragraph(f"Key differences: {profile.key_differences}")
                if profile.evidence:
                    doc.add_paragraph(f"Evidence: {profile.evidence}")
        else:
            doc.add_paragraph("No competitors identified.")

        comparison_rows = [r for r in (coerce_comparison_row(r) for r in (paper.get("comparison_table") or [])) if r]
        add_heading("Comparison Table")
        if comparison_rows:
            # Column order: the evaluated vendor first, then every other
            # named competitor column encountered, in first-seen order.
            columns: list[str] = [vendor_name] if any(vendor_name in r.values for r in comparison_rows) else []
            for r in comparison_rows:
                for key in r.values:
                    if key not in columns:
                        columns.append(key)
            if not columns:
                columns = [vendor_name]

            table = doc.add_table(rows=1, cols=len(columns) + 2)
            table.autofit = True
            headers = ["Criterion", *columns, "Notes / Considerations"]
            for i, htext in enumerate(headers):
                set_cell(table.rows[0].cells[i], htext, header=True)
            for row_model in comparison_rows:
                row = table.add_row()
                set_cell(row.cells[0], row_model.criterion)
                for i, col in enumerate(columns, start=1):
                    set_cell(row.cells[i], row_model.values.get(col, ""))
                set_cell(row.cells[len(columns) + 1], row_model.notes)
        else:
            doc.add_paragraph("No comparison data available.")

        add_heading("Visual Summary")
        architecture_overview = paper.get("architecture_overview", "")
        if architecture_overview:
            doc.add_paragraph(architecture_overview)
        key_insights = paper.get("key_insights") or paper.get("visual_summary") or []
        if key_insights:
            add_heading("Key Insights & Implications", level=2)
            add_bullets(key_insights)

        add_heading("Analyst Take / Recommendation")
        add_body(paper.get("recommendation", ""))

        add_heading("Additional Resources / References")
        references = paper.get("references") or []
        if references:
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            set_cell(table.rows[0].cells[0], "Source Type", header=True)
            set_cell(table.rows[0].cells[1], "Link / Reference", header=True)
            for ref in references:
                row = table.add_row()
                if isinstance(ref, dict):
                    set_cell(row.cells[0], str(ref.get("type") or ref.get("label") or "Reference"))
                    set_cell(row.cells[1], str(ref.get("url", "")))
                else:
                    set_cell(row.cells[0], "Reference")
                    set_cell(row.cells[1], str(ref))
        else:
            doc.add_paragraph("No references captured.")

        doc.save(str(out_path))
        return out_path
